from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_default_style(document):
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)


def set_page_layout(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_title_page(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROJECT REPORT")
    run.bold = True
    run.font.size = Pt(24)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Job and Internship Application Tracking Pipeline")
    run.bold = True
    run.font.size = Pt(18)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Submitted in partial fulfillment of academic mini-project requirements")

    document.add_paragraph()
    lines = [
        "Student Name: ______________________________",
        "Roll Number: _______________________________",
        "Course / Program: ___________________________",
        "Department: ________________________________",
        "Institute: _________________________________",
        f"Submission Date: {date.today().isoformat()}",
        "Guide / Faculty Name: _______________________",
    ]
    for line in lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)

    document.add_page_break()


def add_heading(document, text, level=1):
    document.add_heading(text, level=level)


def add_paras(document, paragraphs):
    for text in paragraphs:
        document.add_paragraph(text)


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_code_block(document, title, code_text):
    if title:
        p = document.add_paragraph()
        r = p.add_run(title)
        r.bold = True
    code_par = document.add_paragraph()
    run = code_par.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def chapter_preliminaries(document):
    add_heading(document, "Certificate", 1)
    add_paras(document, [
        "This is to certify that the project titled 'Job and Internship Application Tracking Pipeline' is a bona fide work carried out by the student as part of the mini-project requirement. The work presented here is original and has been completed under academic guidance.",
        "Faculty Signature: ____________________",
        "Date: ____________________",
    ])

    add_heading(document, "Declaration", 1)
    add_paras(document, [
        "I hereby declare that this report and the project implementation are my original work and have not been submitted elsewhere for any degree or diploma. All references used in this work have been acknowledged.",
        "Student Signature: ____________________",
        "Date: ____________________",
    ])

    add_heading(document, "Acknowledgement", 1)
    add_paras(document, [
        "I express my sincere gratitude to my faculty guide and department for their continuous support and valuable feedback. I also acknowledge publicly available documentation of JavaScript, Firebase, Chart.js, and XLSX libraries that helped in implementing and validating this project.",
    ])

    add_heading(document, "Abstract", 1)
    add_paras(document, [
        "This project presents a dynamic web application for tracking job and internship applications using HTML, CSS, and JavaScript. The system organizes applications into a six-stage Kanban pipeline, supports search and filtering, and provides analytics through charts and KPI summaries. The application includes dual persistence modes: a local guest mode using browser localStorage and an authenticated cloud mode using Firebase Auth and Firestore.",
        "The solution demonstrates practical use of DOM manipulation, event-driven programming, client-side data processing, and user interface state synchronization. Import and export support for CSV/Excel improves usability for real-world job hunt workflows. The final implementation is tested through runtime smoke checks and diagnostic validation and is documented with code-level explanations of key dynamic modules.",
    ])

    add_heading(document, "Keywords", 2)
    add_bullets(document, [
        "Kanban workflow",
        "DOM manipulation",
        "Event handling",
        "Firebase Auth and Firestore",
        "Analytics dashboard",
        "CSV/Excel import-export",
    ])

    document.add_page_break()


def chapter_introduction(document):
    add_heading(document, "Chapter 1: Introduction", 1)
    add_heading(document, "1.1 Background", 2)
    add_paras(document, [
        "Students applying to internships and jobs usually maintain data in fragmented files, making it difficult to monitor progress and plan follow-ups. A stage-based tracker can reduce this fragmentation by combining records, workflow transitions, and metrics in one interface.",
    ])

    add_heading(document, "1.2 Problem Statement", 2)
    add_paras(document, [
        "There is a need for a lightweight dynamic web application that allows users to create, update, classify, and analyze job application entries without requiring a complex backend setup.",
    ])

    add_heading(document, "1.3 Objective", 2)
    add_bullets(document, [
        "Develop a small but fully functional dynamic mini-application.",
        "Demonstrate DOM manipulation and JavaScript event handling.",
        "Provide interactive workflow features beyond static forms.",
        "Generate clear analytics and reporting outputs.",
        "Document all major features with code explanations.",
    ])

    add_heading(document, "1.4 Scope", 2)
    add_bullets(document, [
        "In scope: CRUD operations, drag-and-drop stage movement, filtering, analytics, CSV/Excel import-export, theme support.",
        "Out of scope: multi-user collaboration, custom backend API server, automated email reminders.",
    ])

    add_heading(document, "1.5 Report Organization", 2)
    add_paras(document, [
        "This report is organized into architecture, implementation details, feature-wise code explanations, testing and evaluation, and appendices for screenshots and additional artifacts.",
    ])


def chapter_requirements_design(document):
    add_heading(document, "Chapter 2: Requirement Analysis and Design", 1)
    add_heading(document, "2.1 Functional Requirements", 2)
    add_numbered(document, [
        "User can sign in with email/password when Firebase is configured.",
        "User can continue in guest mode without cloud setup.",
        "User can add, edit, delete, and move applications across stages.",
        "User can search and filter by stage, company, and keywords.",
        "User can access analytics with charts and KPI cards.",
        "User can import CSV/Excel and export CSV reports.",
    ])

    add_heading(document, "2.2 Non-Functional Requirements", 2)
    add_bullets(document, [
        "Usability: clear UI with stage visibility and status counters.",
        "Reliability: app must work in local guest mode without Firebase.",
        "Performance: filtering and rendering should feel responsive for normal dataset sizes.",
        "Security: credentials must stay outside tracked source via local config strategy.",
        "Maintainability: modular JavaScript files for auth, core app logic, and Firebase wrappers.",
    ])

    add_heading(document, "2.3 Tools and Technologies", 2)
    table = document.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Technology"
    hdr[2].text = "Purpose"
    rows = [
        ("Frontend", "HTML5, CSS3", "Page structure and visual styling"),
        ("Logic", "Vanilla JavaScript", "State, events, rendering"),
        ("Analytics", "Chart.js", "Bar and doughnut charts"),
        ("Import", "XLSX", "Excel parsing in browser"),
        ("Cloud", "Firebase Auth + Firestore", "User auth and persistent storage"),
        ("Optional utility", "Python + Playwright", "Automated table extraction script"),
    ]
    for layer, tech, purpose in rows:
        cells = table.add_row().cells
        cells[0].text = layer
        cells[1].text = tech
        cells[2].text = purpose

    add_heading(document, "2.4 System Architecture", 2)
    add_bullets(document, [
        "index.html: contains auth screen, dashboard, analytics view, and utility controls.",
        "js/auth.js: initializes auth flow, guest mode, and login event handling.",
        "js/app.js: handles main state, Kanban rendering, analytics filters, and import/export.",
        "js/firebase-config.js: wraps Firebase initialization, auth API, and Firestore operations.",
    ])


def chapter_implementation(document):
    add_heading(document, "Chapter 3: Implementation and Dynamic Features", 1)
    add_heading(document, "3.1 Core Dynamic Features", 2)
    add_bullets(document, [
        "Interactive form-driven CRUD operations",
        "Drag-and-drop stage transitions",
        "Dynamic filtering and sorting",
        "Analytics chart generation from live state",
        "Theme synchronization across auth and app screens",
        "Import and export data pipelines",
    ])

    add_heading(document, "3.2 Feature 1 - State and Persistence", 2)
    add_paras(document, [
        "The project uses a centralized state object for application entries, filters, and theme. Local storage wrappers are used to protect against restricted browser contexts. This keeps the app functional even when persistence is unavailable.",
    ])

    code_1 = """const state = {
  applications: [],
  currentWeeklyGoal: parseInt(safeStorageGet(GOAL_KEY), 10) || 10,
  currentFilter: \"all\",
  currentCompanyFilter: \"all\",
  currentSort: \"deadline-asc\",
  searchQuery: \"\",
  editingId: null,
  theme: safeStorageGet(THEME_KEY) || \"light\"
};

function safeStorageSet(key, value) {
  try {
    localStorage.setItem(key, value);
  } catch {
    // Storage may be blocked; keep app functional without persistence.
  }
}"""
    add_code_block(document, "Code Snippet A (js/app.js): state and safe persistence wrappers", code_1)

    add_heading(document, "3.3 Feature 2 - Add/Edit/Delete Application (DOM + Events)", 2)
    add_paras(document, [
        "The save handler reads form fields, infers missing values from a job link, validates mandatory information, and then either updates an existing record or inserts a new one. Changes are persisted to Firebase when available, otherwise to localStorage.",
    ])

    code_2 = """function saveApplicationFromForm(event) {
  event.preventDefault();
  const id = getValue(\"form-id\");
  const typedCompany = getValue(\"form-company\").trim();
  const typedRole = getValue(\"form-role\").trim();
  const link = getValue(\"form-link\").trim();

  const inferred = inferFromJobLink(link);
  const company = typedCompany || (inferred ? inferred.company : \"\") || (link ? \"Unknown Company\" : \"\");
  const role = typedRole || (inferred ? inferred.role : \"\") || (link ? \"Unknown Role\" : \"\");

  if (!company || !role) {
    showToast(\"Please provide a Job Link or upload an attachment\", \"error\");
    return;
  }

  // update state and persist to Firebase or localStorage
}"""
    add_code_block(document, "Code Snippet B (js/app.js): form-driven save operation", code_2)

    add_heading(document, "3.4 Feature 3 - Drag-and-Drop Kanban Transition", 2)
    add_paras(document, [
        "Kanban columns register drag-over and drop events. Dropped items trigger stage updates, persistence writes, and immediate UI rerender. This demonstrates event handling and dynamic DOM updates.",
    ])

    code_3 = """function bindKanbanDnD() {
  document.querySelectorAll(\".column-list\").forEach((column) => {
    column.addEventListener(\"dragover\", (event) => {
      event.preventDefault();
      event.dataTransfer.dropEffect = \"move\";
      column.classList.add(\"drag-over\");
    });

    column.addEventListener(\"drop\", (event) => {
      event.preventDefault();
      column.classList.remove(\"drag-over\");
      const appId = event.dataTransfer.getData(\"text/plain\");
      const targetStage = column.dataset.stage;
      moveApplicationToStage(appId, targetStage);
    });
  });
}"""
    add_code_block(document, "Code Snippet C (js/app.js): drag-and-drop binding", code_3)

    add_heading(document, "3.5 Feature 4 - Analytics Filter Pipeline", 2)
    add_paras(document, [
        "Analytics applies composable filters (status, date range, company, location, keyword, free-text). Filtered rows are sorted and used to update counters, chips, table, and charts in a single flow.",
    ])

    code_4 = """function an_applyFilters() {
  let rows = state.applications.filter((app) => {
    const code = an_stageCode(app.stage);
    if (!anState.statuses.has(code)) return false;
    // date, company, location, keyword, and search checks
    return true;
  });

  rows = an_sortRows(rows, anState.sort);
  anState.filtered = rows;
  an_renderStatusMeta(rows);
  an_renderActiveChips();
  an_updateStats(rows);
  an_renderTable(rows);
  an_buildCharts(rows);
}"""
    add_code_block(document, "Code Snippet D (js/app.js): analytics filtering and rendering pipeline", code_4)

    add_heading(document, "3.6 Feature 5 - Dynamic Chart Rendering", 2)
    add_paras(document, [
        "The chart module destroys prior chart instances and creates fresh chart datasets from filtered rows. The doughnut chart also adapts border color based on theme state, showing data-driven plus theme-driven rendering behavior.",
    ])

    code_5 = """function an_buildCharts(rows = anState.filtered) {
  if (!anState.isOpen || typeof Chart === \"undefined\") return;

  if (anCharts.bar) anCharts.bar.destroy();
  if (anCharts.donut) anCharts.donut.destroy();

  // Build bar chart and doughnut chart from stage counts
}"""
    add_code_block(document, "Code Snippet E (js/app.js): chart lifecycle management", code_5)

    add_heading(document, "3.7 Feature 6 - Import and Export", 2)
    add_paras(document, [
        "The import pipeline supports CSV and Excel formats, normalizes columns, infers missing details from links, validates entries, and persists each imported record. Export generates CSV rows from current app state or filtered analytics state.",
    ])

    code_6 = """async function onImportFileChange(event) {
  const [file] = event.target.files || [];
  if (!file) return;

  const imported = await parseImportFile(file);
  imported.forEach((entry) => {
    const normalized = normalizeImportedEntry(entry);
    // map, validate, and save
  });
}

function exportToCSV() {
  if (!state.applications.length) {
    showToast(\"No applications to export\", \"warning\");
    return;
  }
  // build csv string and trigger download
}"""
    add_code_block(document, "Code Snippet F (js/app.js): import/export flow", code_6)


def chapter_auth_cloud(document):
    add_heading(document, "Chapter 4: Authentication and Cloud Data Flow", 1)
    add_heading(document, "4.1 Auth Module", 2)
    add_paras(document, [
        "The authentication layer validates user credentials and toggles between auth screen and app layout. It also supports guest mode for no-cloud usage and binds theme synchronization on both screens.",
    ])

    code_7 = """function bindLoginForm() {
  const form = getEl('auth-login-form');
  form.addEventListener('submit', async (event) => {
    event.preventDefault();
    const email = (getEl('auth-email-input')?.value || '').trim();
    const password = getEl('auth-password-input')?.value || '';
    await FirebaseAPI.auth.signInWithEmail(email, password);
  });
}

function startGuestMode() {
  isGuestMode = true;
  currentUser = { uid: 'guest-local', isGuest: true, email: 'Guest mode' };
  showAppLayout();
}"""
    add_code_block(document, "Code Snippet G (js/auth.js): login submit and guest mode", code_7)

    add_heading(document, "4.2 Firebase Wrapper Module", 2)
    add_paras(document, [
        "The Firebase configuration module isolates cloud operations from UI code. It exports a narrow API for sign-in, sign-out, CRUD operations, and settings persistence.",
    ])

    code_8 = """window.FirebaseAPI = {
  initialize: initializeFirebase,
  isReady: () => !!auth && !!db,
  auth: {
    signInWithEmail,
    signOut,
    getCurrentUser,
    onAuthStateChanged
  },
  db: {
    saveApplication: saveApplicationToFirestore,
    loadApplications: loadApplicationsFromFirestore,
    listenToApplications: listenToApplications,
    deleteApplication: deleteApplicationFromFirestore,
    saveSettings: saveUserSettings,
    loadSettings: loadUserSettings
  }
};"""
    add_code_block(document, "Code Snippet H (js/firebase-config.js): exported API boundary", code_8)

    add_heading(document, "4.3 Security Approach", 2)
    add_bullets(document, [
        "Public repository keeps placeholders in js/firebase-config.js.",
        "Real credentials are stored locally in js/firebase-config.local.js (gitignored).",
        "Firestore rules should enforce user-scoped access by UID.",
    ])


def chapter_results_testing(document):
    add_heading(document, "Chapter 5: Results, Validation, and Evaluation", 1)
    add_heading(document, "5.1 Execution Results", 2)
    add_bullets(document, [
        "Application served successfully over localhost during smoke testing.",
        "UI screens (auth, dashboard, analytics) rendered and switched correctly.",
        "CRUD, stage movement, and import/export flows behaved as expected.",
        "Analytics charts updated according to active filters.",
    ])

    add_heading(document, "5.2 Evaluation Criteria Mapping", 2)
    table = document.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Criterion"
    hdr[1].text = "Weight"
    hdr[2].text = "Project Evidence"
    rows = [
        ("Functionality", "40%", "Kanban CRUD, drag-drop, analytics, import/export"),
        ("Code Quality", "30%", "Modular JS files, helper functions, validation checks"),
        ("Documentation", "20%", "Structured report with code snippets and explanations"),
        ("Creativity & Completeness", "10%", "Dual mode persistence, analytics dashboard, themed UI"),
    ]
    for c1, c2, c3 in rows:
        cells = table.add_row().cells
        cells[0].text = c1
        cells[1].text = c2
        cells[2].text = c3

    add_heading(document, "5.3 Test Cases", 2)
    test_table = document.add_table(rows=1, cols=4)
    h = test_table.rows[0].cells
    h[0].text = "Test ID"
    h[1].text = "Scenario"
    h[2].text = "Expected"
    h[3].text = "Status"

    tests = [
        ("TC-01", "Add new application", "Record appears in selected stage", "Pass"),
        ("TC-02", "Edit existing application", "Changes reflected after save", "Pass"),
        ("TC-03", "Delete application", "Record removed after confirmation", "Pass"),
        ("TC-04", "Drag card to new stage", "Stage updates and UI rerenders", "Pass"),
        ("TC-05", "Filter analytics by status/date", "Rows and KPIs update", "Pass"),
        ("TC-06", "Export CSV", "Download contains current records", "Pass"),
        ("TC-07", "Import CSV/Excel", "Valid rows mapped and inserted", "Pass"),
        ("TC-08", "Guest mode", "Data persists in localStorage", "Pass"),
    ]
    for tid, scenario, expected, status in tests:
        cells = test_table.add_row().cells
        cells[0].text = tid
        cells[1].text = scenario
        cells[2].text = expected
        cells[3].text = status


def chapter_reflection(document):
    add_heading(document, "Chapter 6: Reflection and Lessons Learned", 1)
    add_heading(document, "6.1 Challenges Faced", 2)
    add_bullets(document, [
        "Maintaining consistency between dashboard and analytics view states.",
        "Handling both local and Firebase persistence paths cleanly.",
        "Normalizing inconsistent CSV/Excel column names during import.",
        "Keeping UI responsive while updating multiple sections after each state change.",
    ])

    add_heading(document, "6.2 Solutions Applied", 2)
    add_bullets(document, [
        "Centralized state object and deterministic render flow.",
        "Wrapper-based persistence APIs to isolate storage differences.",
        "Normalization helpers for robust import parsing.",
        "Modular functions for filters, rendering, and chart updates.",
    ])

    add_heading(document, "6.3 Learning Outcomes", 2)
    add_bullets(document, [
        "Practical DOM and event-driven design in Vanilla JavaScript.",
        "Importance of clean module boundaries in frontend architecture.",
        "Handling edge cases in browser storage and external file parsing.",
        "Writing maintainable project documentation aligned with evaluation rubrics.",
    ])

    add_heading(document, "6.4 Future Enhancements", 2)
    add_bullets(document, [
        "Add automated tests (unit and end-to-end).",
        "Introduce reminder notifications for deadlines and follow-ups.",
        "Improve analytics trend history and comparative time windows.",
        "Refactor large CSS override files into a unified design system.",
    ])


def chapter_conclusion(document):
    add_heading(document, "Chapter 7: Conclusion", 1)
    add_paras(document, [
        "The Job and Internship Application Tracking Pipeline fulfills the mini-project objective of building a dynamic and functional web application using HTML, CSS, and JavaScript. The project includes interactive event handling, practical data operations, and measurable analytics outputs. It demonstrates both technical implementation and report-level explanation required for academic evaluation.",
    ])


def chapter_appendices(document):
    add_heading(document, "Appendix A: Screenshot Placeholders", 1)
    add_numbered(document, [
        "Login screen with email/password and guest mode buttons.",
        "Dashboard showing stage columns and application cards.",
        "Add/Edit modal form and validation message.",
        "Drag-and-drop action from one stage to another.",
        "Analytics view with active filters and KPI cards.",
        "Bar and doughnut chart outputs.",
        "CSV import success message and updated rows.",
        "CSV export file opened in spreadsheet view.",
    ])

    add_heading(document, "Appendix B: File Structure", 1)
    add_code_block(document, "", """.
|-- index.html
|-- README.md
|-- css/
|   |-- styles.css
|   |-- redesign.css
|   `-- auth-screen-styles.css
|-- js/
|   |-- app.js
|   |-- auth.js
|   |-- firebase-config.js
|   |-- firebase-config.example.js
|   `-- firebase-config.local.js
|-- docs/
|   `-- PROJECT_DOCUMENTATION.md
|-- data/
|   `-- sample_applications.csv
`-- tools/
    |-- sap_web_scraping.py
    `-- generate_academic_report.py
""")

    add_heading(document, "Appendix C: Reference URLs", 1)
    add_bullets(document, [
        "MDN Web Docs: DOM APIs and JavaScript language references",
        "Firebase Documentation: Auth and Firestore",
        "Chart.js Documentation",
        "SheetJS (XLSX) Documentation",
    ])


def main():
    document = Document()
    set_default_style(document)
    set_page_layout(document)

    add_title_page(document)
    chapter_preliminaries(document)
    chapter_introduction(document)
    chapter_requirements_design(document)
    chapter_implementation(document)
    chapter_auth_cloud(document)
    chapter_results_testing(document)
    chapter_reflection(document)
    chapter_conclusion(document)
    chapter_appendices(document)

    output_path = "docs/Project_Report_Academic.docx"
    document.save(output_path)
    print(f"Report generated: {output_path}")


if __name__ == "__main__":
    main()
